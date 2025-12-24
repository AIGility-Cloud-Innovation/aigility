# [服务层] 对外暴露的统一入口 (RAGService)
"""
RAG Service - 检索增强生成服务

使用方式:
    from aigility.rag import RAGService, RAGConfig, EmbeddingConfig, VectorStoreConfig
    
    # 配置由业务项目传入
    config = RAGConfig(
        embedding=EmbeddingConfig(provider="dashscope", api_key="your-key"),
        vector_store=VectorStoreConfig(provider="chroma", persist_path="./my_db")
    )
    service = RAGService(config=config)
"""

import os
import shutil
import hashlib
from typing import Optional, Dict, List

# 抑制 tokenizers 并行警告
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")

from .config import RAGConfig
from .embeddings.factory import EmbeddingFactory
from .vector_stores.factory import VectorStoreFactory
from .ingestion import IngestionManager

# 可选依赖，延迟导入
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    import jieba
    HAS_NLP_DEPS = True
except ImportError:
    HAS_NLP_DEPS = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


class RAGService:
    """RAG 服务主类，提供文档入库和检索能力"""
    
    def __init__(self, config: Optional[RAGConfig] = None):
        """
        初始化 RAG 服务
        
        Args:
            config: RAG 配置对象，包含 embedding、vector_store、ingestion 配置。
                    如果不传，使用默认配置（HuggingFace + Chroma 本地模式）
        """
        self.config = config or RAGConfig()
        
        print(f"🔧 Initializing RAG with: Embedding={self.config.embedding.provider}, Store={self.config.vector_store.provider}")

        # 1. 工厂生产 Embedding
        self.embedding_model = EmbeddingFactory.get_embedding_model(self.config.embedding)
        
        # 2. 工厂生产 Vector Store (注入 embedding)
        self.vector_store = VectorStoreFactory.get_vector_store(
            self.config.vector_store, 
            self.embedding_model
        )
        
        # 3. 初始化数据处理模块
        self.ingestion = IngestionManager(self.config.ingestion)
        
        # 文档元信息存储
        self.doc_meta_info: Dict[str, dict] = {}
        self.global_doc_keywords: List[str] = []

    def _read_file(self, file_path: str) -> str:
        """
        统一读取 txt/docx/pdf 格式文件，提取纯文本内容
        
        Args:
            file_path: 文件绝对路径
            
        Returns:
            清洗后的纯文本字符串（失败返回空字符串）
        """
        file_suffix = os.path.splitext(file_path)[-1].lower()
        pure_text = ""

        try:
            if file_suffix == ".txt":
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        pure_text = f.read()
                except UnicodeDecodeError:
                    with open(file_path, "r", encoding="gbk") as f:
                        pure_text = f.read()

            elif file_suffix == ".docx":
                if not HAS_DOCX:
                    raise ImportError("请安装 python-docx: pip install python-docx")
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        pure_text += para_text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = "\t".join([cell.text.strip() for cell in row.cells])
                        if row_text:
                            pure_text += row_text + "\n"

            elif file_suffix == ".pdf":
                if not HAS_PDF:
                    raise ImportError("请安装 pdfplumber: pip install pdfplumber")
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pure_text += page_text.strip() + "\n"

            else:
                raise ValueError(f"不支持的文件格式：{file_suffix}，仅支持 txt/docx/pdf")

            # 文本清洗
            pure_text = pure_text.replace("\r", "").replace("\t", " ").replace("  ", " ").strip()
            return pure_text

        except Exception as e:
            print(f"❌ 读取文件 {file_path} 失败：{str(e)}")
            return ""

    def _extract_doc_meta(self, doc_text: str, doc_name: str) -> dict:
        """
        从纯文本中提取文档元信息（关键词+摘要）
        
        Args:
            doc_text: 文档纯文本内容
            doc_name: 文档文件名
            
        Returns:
            包含关键词和摘要的元信息字典
        """
        # 生成文档摘要
        doc_summary = doc_text[:200].strip() if len(doc_text) > 200 else doc_text.strip()

        # 生成核心关键词
        doc_keywords = []
        if HAS_NLP_DEPS:
            try:
                word_list = jieba.lcut(doc_text.replace("\n", "").strip())
                seg_text = " ".join(word_list)
                corpus = [seg_text]

                tfidf = TfidfVectorizer(max_features=10, stop_words=None)
                tfidf.fit_transform(corpus)
                doc_keywords = tfidf.get_feature_names_out().tolist()
            except Exception as e:
                print(f"⚠️ 提取 {doc_name} 关键词失败：{str(e)}，使用文件名兜底")
                doc_keywords = [word for word in doc_name.split(".")[0].split("_") if word.strip()]
        else:
            # 无 NLP 依赖时，使用文件名作为关键词
            doc_keywords = [word for word in doc_name.split(".")[0].split("_") if word.strip()]

        doc_meta = {
            "name": doc_name,
            "keywords": doc_keywords,
            "summary": doc_summary
        }

        # 更新全局元信息
        self.doc_meta_info[doc_name] = doc_meta
        self.global_doc_keywords.extend(doc_keywords)
        self.global_doc_keywords = list(set(self.global_doc_keywords))

        return doc_meta

    def add_file(self, file_path: str):
        """
        加载文件 -> 切分 -> 存入向量库（带去重）
        
        Args:
            file_path: 文件路径（支持相对路径和绝对路径）
        """
        try:
            if not isinstance(file_path, str):
                raise TypeError(f"file_path must be str, got {type(file_path)}")
            
            # 转换为绝对路径
            if not os.path.isabs(file_path):
                file_path = os.path.abspath(file_path)
            
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            # 计算文件 hash，用于去重
            with open(file_path, "rb") as f:
                file_hash = hashlib.md5(f.read()).hexdigest()
            
            doc_name = os.path.basename(file_path)
            
            # 检查是否已添加过该文件
            existing = self.vector_store.get(where={"file_hash": file_hash}, limit=1)
            if existing and existing.get("ids"):
                print(f"⚠️ 文件已存在，跳过添加: {file_path}")
                if doc_name not in self.doc_meta_info:
                    doc_text = self._read_file(file_path)
                    if doc_text:
                        self._extract_doc_meta(doc_text, doc_name)
                        print(f"✅ 已加载 {doc_name} 的元信息")
                return
            
            print(f"📄 Processing file: {file_path}")
            
            # 读取文件内容
            doc_text = self._read_file(file_path)
            if not doc_text:
                print(f"❌ {doc_name} 提取纯文本为空，无法添加到知识库")
                return
            
            # 生成并存储元信息
            doc_meta = self._extract_doc_meta(doc_text, doc_name)
            print(f"✅ 成功生成 {doc_name} 元信息：关键词={doc_meta['keywords']}")
            
            raw_docs = self.ingestion.load_file(file_path)
            chunks = self.ingestion.process_raw_docs(raw_docs, file_path)
            
            # 给每个 chunk 添加 file_hash 元数据
            for chunk in chunks:
                chunk.metadata["file_hash"] = file_hash
            
            if chunks:
                self.vector_store.add_documents(chunks)
                print(f"✅ Successfully added {len(chunks)} chunks to {self.config.vector_store.provider}.")
            else:
                print("⚠️ No content found in file.")
                
        except Exception as e:
            print(f"❌ Error adding file {file_path}: {str(e)}")
            raise e

    def search(self, query: str) -> str:
        """
        检索相关文档
        
        Args:
            query: 检索查询语句
            
        Returns:
            格式化的检索结果字符串
        """
        try:
            docs = self.vector_store.similarity_search(query, k=self.config.search_top_k)
            
            if not docs:
                return ""

            results = []
            for doc in docs:
                source = doc.metadata.get("source", "Unknown")
                content = doc.page_content.replace("\n", " ")
                results.append(f"Source: {source}\nContent: {content}")
            
            return "\n\n".join(results)
            
        except Exception as e:
            print(f"❌ Search failed: {str(e)}")
            return ""

    def clear_knowledge_base(self):
        """(危险操作) 清空知识库"""
        try:
            if self.config.vector_store.provider == "chroma":
                if os.path.exists(self.config.vector_store.persist_path):
                    shutil.rmtree(self.config.vector_store.persist_path)
                    os.makedirs(self.config.vector_store.persist_path, exist_ok=True)
                    print(f"✅ Chroma 知识库 {self.config.vector_store.persist_path} 已清空")
                    
            elif self.config.vector_store.provider == "milvus":
                from pymilvus import utility
                if utility.has_collection(self.config.vector_store.collection_name):
                    utility.drop_collection(self.config.vector_store.collection_name)
                    print(f"✅ Milvus 集合 {self.config.vector_store.collection_name} 已删除")
                    
            elif self.config.vector_store.provider == "faiss":
                faiss_index_path = os.path.join(
                    self.config.vector_store.persist_path,
                    f"{self.config.vector_store.collection_name}.index"
                )
                if os.path.exists(faiss_index_path):
                    os.remove(faiss_index_path)
                if os.path.exists(self.config.vector_store.persist_path):
                    shutil.rmtree(self.config.vector_store.persist_path)
                    os.makedirs(self.config.vector_store.persist_path, exist_ok=True)
                print(f"✅ FAISS 索引 {faiss_index_path} 已清空")
            else:
                print(f"⚠️ 不支持清空 {self.config.vector_store.provider} 知识库")
                
        except Exception as e:
            print(f"❌ 清空知识库失败: {str(e)}")

    def get_global_keywords(self) -> List[str]:
        """获取全局文档关键词（供 rag_decision 节点使用）"""
        return self.global_doc_keywords

    def get_doc_meta(self, doc_name: str) -> dict:
        """获取单个文档元信息"""
        return self.doc_meta_info.get(doc_name, {})

    def get_all_doc_meta(self) -> Dict[str, dict]:
        """获取所有文档元信息"""
        return self.doc_meta_info


__all__ = ["RAGService"]
