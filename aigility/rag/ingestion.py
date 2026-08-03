# ingestion.py
"""
轻量级文档解析与切分管理器 (RAGFlow 逻辑优化版)

核心改进:
1. PDF: 使用 pdfplumber 基于坐标还原布局，支持表格转 Markdown。
2. Excel: 借鉴 RAGFlow 将行数据序列化为 "键:值" 格式，保留语义。
3. Docx: 区分处理文本段落与表格，防止表格内容错乱。
4. Markdown: 支持基于 AST 的智能切分，保持标题-内容关联性。
"""

import os
import re
import logging
from typing import List, Dict, Any, Optional
from hashlib import md5
import pandas as pd

# 轻量级依赖
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .config import IngestionConfig

# 尝试导入 Markdown AST 切分器
try:
    from .markdown_splitter import MarkdownASTSplitter
    MARKDOWN_AST_AVAILABLE = True
except ImportError:
    MARKDOWN_AST_AVAILABLE = False

class IngestionManager:
    """通用文档解析与切分管理器"""

    def __init__(self, config: IngestionConfig):
        self.config = config
        self._fallback_splitter = None

        self.duplicate_cache: Dict[str, set] = {}

        # 默认使用 AST 切分器（如果可用）
        if MARKDOWN_AST_AVAILABLE:
            self._ast_splitter = MarkdownASTSplitter(
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap,
                context_buffer_size=config.context_buffer_size,
                min_chunk_size=config.min_chunk_length,
                enable_small2big=config.enable_small2big,
                parent_chunk_size=config.parent_chunk_size,
                child_chunk_size=config.child_chunk_size,
                child_chunk_overlap=config.child_chunk_overlap,
            )
            self._use_ast = True
            print("✅ 使用 Markdown AST 切分器（保持标题-内容关联）")
        else:
            self._ast_splitter = None
            self._use_ast = False
            print("⚠️ markdown-it-py 未安装，降级使用传统切分器")

    @property
    def splitter(self):
        """获取兜底的传统切分器"""
        if self._fallback_splitter is None:
            self._fallback_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.config.chunk_size,
                chunk_overlap=self.config.chunk_overlap,
                separators=[
                    r"\n#{1,6} ",
                    r"\n\*\*\*+\n",
                    r"\n---+\n",
                    r"\n___+\n",
                    r"\n\n",
                    r"\n",
                    "。", "！", "？",
                    "; ", "；",
                    " ", ""
                ],
                is_separator_regex=True,
                length_function=len
            )
        return self._fallback_splitter

    def load_file(self, file_path: str) -> List[Document]:
        """统一入口：根据文件类型分发处理逻辑"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = os.path.splitext(file_path)[-1].lower()
        file_name = os.path.basename(file_path)
        
        try:
            if ext == ".pdf":
                texts = self._parse_pdf_smart(file_path)
            elif ext in [".xlsx", ".xls", ".csv"]:
                texts = self._parse_excel_semantic(file_path, ext)
            elif ext in [".docx", ".doc"]:
                texts = self._parse_docx_struct(file_path)
            elif ext in [".txt", ".md"]:
                texts = self._parse_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")
            
            # 封装为 Document 对象
            docs = []
            for i, text in enumerate(texts):
                if not text.strip():
                    continue
                docs.append(Document(
                    page_content=text,
                    metadata={
                    "file_path": file_path,
                        "file_name": file_name,
                    "file_type": ext,
                        "chunk_index": i
                    }
                ))
            return docs
        except Exception as e:
            logging.error(f"Error loading {file_path}: {e}")
            return []

    def _parse_pdf_smart(self, file_path: str) -> List[str]:
        """
        [借鉴 RAGFlow] PDF 解析逻辑
        策略: 使用 pdfplumber 获取布局信息。
        1. 表格: 提取并转为 Markdown 格式，保留结构。
        2. 文本: 尝试基于坐标聚类 (虽然不如 DeepDoc 准确，但比 pypdf 强)。
        """
        if not HAS_PDFPLUMBER:
            logging.warning("pdfplumber not installed, falling back to basic PDF parsing")
            return self._parse_pdf_basic(file_path)

        full_text_pages = []

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_content = []
                
                # 1. 提取表格 (优先提取，并在原文中占位)
                # pdfplumber 的 extract_tables 基于线条检测，比 DeepDoc 的视觉模型轻量
                tables = page.extract_tables()
                table_texts = []
                if tables:
                    for table in tables:
                        # 将表格转为 Markdown 格式
                        # 过滤 None 值
                        clean_table = [[str(cell or "").replace("\n", " ") for cell in row] for row in table]
                        if clean_table and len(clean_table) > 1:
                            try:
                                header = clean_table[0]
                                data_rows = clean_table[1:]
                                # 手动生成标准 Markdown 表格（不依赖 tabulate）
                                header_line = "| " + " | ".join(header) + " |"
                                separator = "| " + " | ".join(["---"] * len(header)) + " |"
                                data_lines = ["| " + " | ".join(row) + " |" for row in data_rows]
                                md_table = "\n".join([header_line, separator] + data_lines)
                                table_texts.append(f"\n\n{md_table}\n\n")
                            except Exception:
                                pass # 表格结构过于复杂，跳过
                
                # 2. 提取文本
                text = page.extract_text()
                if text:
                    # 【关键优化】移除页码标记（常见格式）
                    # 格式: "第 X 页", "Page X", "- X -", "X / Y" 等
                    text = re.sub(r'第\s*\d+\s*页', '', text)
                    text = re.sub(r'[Pp]age\s*\d+', '', text)
                    text = re.sub(r'-\s*\d+\s*-', '', text)
                    text = re.sub(r'\d+\s*/\s*\d+', '', text)
                    # 移除连续的空行
                    text = re.sub(r'\n{3,}', '\n\n', text)
                    
                    if text.strip():
                        page_content.append(text)
                
                # 合并内容: 文本 + 表格
                # 注意: 理想情况是将表格插回原文位置，pdfplumber 可以做到但代码量大
                # 这里采用 "文本在前，表格在后" 的策略，保证数据不丢失
                page_str = "\n".join(page_content)
                if table_texts:
                    page_str += "\n".join(table_texts)
                
                full_text_pages.append(page_str)
                
        return full_text_pages

    def _parse_excel_semantic(self, file_path: str, ext: str) -> List[str]:
        """
        [借鉴 RAGFlow - excel_parser.py]
        策略: 将每一行序列化为 "Key: Value" 形式。
        这对 Text-to-SQL 或基于数据的问答至关重要，因为普通的文本切分会把表头和数值切开。
        """
        docs = []
        try:
            if ext == ".csv":
                df = pd.read_csv(file_path)
            else:
                # 仅读取第一个 sheet，或者遍历所有 sheets
                dfs = pd.read_excel(file_path, sheet_name=None)
                # 展平所有 sheet
                df = pd.concat(dfs.values(), ignore_index=True)

            # 数据清洗
            df = df.fillna("")
            
            # 序列化逻辑，参考 excel_parser.py 的 __call__ 方法
            # 格式: "列名1: 值1; 列名2: 值2..."
            text_rows = []
            columns = df.columns.tolist()
            
            for _, row in df.iterrows():
                row_parts = []
                for col in columns:
                    val = str(row[col]).strip()
                    if val:
                        row_parts.append(f"{col}: {val}")
                if row_parts:
                    text_rows.append("; ".join(row_parts))
            
            # Excel 数据通常密集，每 X 行作为一个 chunk，避免 token 溢出
            batch_size = 20 # 假设每行 50 tokens，20 行约 1000 tokens
            for i in range(0, len(text_rows), batch_size):
                batch = text_rows[i:i+batch_size]
                docs.append("\n".join(batch))
                
        except Exception as e:
            logging.error(f"Error parsing Excel: {e}")
            
        return docs

    def _parse_docx_struct(self, file_path: str) -> List[str]:
        """
        [借鉴 RAGFlow - docx_parser.py]
        策略: 区分段落和表格。
        docx_parser.py 会尝试提取表格并线性化。
        """
        doc = DocxDocument(file_path)
        full_text = []
        
        # python-docx 的 iter_block_items 逻辑需要自行实现来保持顺序
        # 这里简化：按 document.element.body 顺序遍历
        
        def iter_block_items(parent):
            from docx.document import Document
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import _Cell, Table
            from docx.text.paragraph import Paragraph

            if isinstance(parent, _Cell):
                # 对于单元格，获取其底层的XML元素
                parent_elm = parent._tc
            else:
                # 对于文档主体，获取其底层XML元素
                # 🔧 修复：在新版python-docx (1.1.2)中，_Body没有element属性，需要使用_element
                parent_elm = parent._body._element if hasattr(parent._body, '_element') else parent._body

            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        for block in iter_block_items(doc):
            # 处理文本段落
            if 'Paragraph' in str(type(block)):
                text = block.text.strip()
                if text:
                    # 优先检测 Word 标题样式，其次基于文本模式识别标题
                    heading_level = self._get_docx_heading_level(block)
                    if not heading_level:
                        heading_level = self._detect_heading_by_text(text)
                    if heading_level:
                        full_text.append(f"{'#' * heading_level} {text}")
                    else:
                        full_text.append(text)
            
            # 处理表格 [关键优化]
            elif 'Table' in str(type(block)):
                # 将表格转为 Markdown 或 线性KV对
                # 这里选择 Markdown，因为 LLM 对 Markdown 表格理解很好
                rows_data = []
                for row in block.rows:
                    cell_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    rows_data.append(cell_data)
                
                if rows_data:
                    try:
                        # 假设第一行是表头
                        header = rows_data[0]
                        data_rows = rows_data[1:]
                        # 手动生成标准 Markdown 表格（不依赖 tabulate）
                        header_line = "| " + " | ".join(header) + " |"
                        separator = "| " + " | ".join(["---"] * len(header)) + " |"
                        data_lines = ["| " + " | ".join(row) + " |" for row in data_rows]
                        md_table = "\n".join([header_line, separator] + data_lines)
                        full_text.append(f"\n{md_table}\n")
                    except Exception:
                        # 如果表格不规范，退化为文本拼接
                        text_table = "\n".join([" | ".join(row) for row in rows_data])
                        full_text.append(f"\n{text_table}\n")

        return ["\n".join(full_text)]

    def _get_docx_heading_level(self, paragraph) -> Optional[int]:
        """
        检测 Word 段落是否为标题样式，返回标题级别

        Args:
            paragraph: python-docx 的 Paragraph 对象

        Returns:
            标题级别 (1-6)，如果不是标题则返回 None
        """
        try:
            style_name = paragraph.style.name.lower() if paragraph.style else ""

            # 直接匹配 "Heading 1" 到 "Heading 6" 样式
            if "heading" in style_name:
                match = re.search(r'heading\s*(\d)', style_name)
                if match:
                    return int(match.group(1))

            # 检查大纲级别 (OutlineLevel)
            pPr = paragraph._p.pPr
            if pPr is not None:
                outlineLvl = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
                if outlineLvl is not None:
                    level = int(outlineLvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '9'))
                    if level < 9:  # 0-8 是有效标题级别
                        return level + 1  # outlineLvl 从 0 开始，标题从 # 开始

            return None
        except Exception:
            return None

    def _detect_heading_by_text(self, text: str) -> Optional[int]:
        """
        基于文本模式识别标题（中英文混合文档）

        Args:
            text: 段落文本

        Returns:
            标题级别 (1-6)，如果不是标题则返回 None
        """
        # 中文数字标题：一、二、三... 十一、十二...
        if re.match(r'^[一二三四五六七八九十]+[、.]', text):
            return 1

        # 数字编号标题：1.1, 7.1, 10.2, 11.2...
        if re.match(r'^\d+\.\d+[\s\.]', text):
            return 2

        # 单层数字标题：1、2、3... 或 1. 2. 3...
        if re.match(r'^\d+[、.]', text):
            return 1

        # 带括号的标题：（一）（二）
        if re.match(r'^[（(][一二三四五六七八九十]+[）)]', text):
            return 2

        return None

    def _parse_txt(self, file_path: str) -> List[str]:
        with open(file_path, 'r', encoding='utf-8') as f:
            return [f.read()]

    def process_documents(self, file_path: str) -> List[Document]:
        """完整流程：加载 -> 合并 -> 清洗 -> 切分 -> 标签"""
        # 1. 智能加载 (替代原来的 load_file)
        raw_docs = self.load_file(file_path)
        
        if not raw_docs:
            return []
        
        # 2. 【关键修复】合并所有页面/段落为一个完整文档
        # 这样切分器的 overlap 才能正常工作（跨页 overlap）
        file_name = raw_docs[0].metadata.get("file_name", "unknown")
        file_type = raw_docs[0].metadata.get("file_type", "unknown")
        
        merged_text = "\n\n".join([doc.page_content for doc in raw_docs])
        merged_doc = Document(
            page_content=merged_text,
            metadata={
                "file_path": file_path,
                "file_name": file_name,
                "file_type": file_type,
                "total_pages": len(raw_docs)
            }
        )
        
        # 3. 文本清洗 (保留 Markdown 表格结构)
        cleaned_docs = self._clean_docs_light([merged_doc])
        
        # 4. 切分 (默认使用 AST 切分器，保持标题-内容关联)
        logging.info(f"🔧 切分配置: chunk_size={self.config.chunk_size}, chunk_overlap={self.config.chunk_overlap}")

        if self._use_ast:
            if self.config.enable_small2big:
                split_docs = self._ast_splitter.split_documents_small2big(cleaned_docs)
            else:
                split_docs = self._ast_splitter.split_documents(cleaned_docs)
        else:
            split_docs = self.splitter.split_documents(cleaned_docs)
        
        # 【调试】打印前几个 chunk 的首尾内容，验证 overlap
        if len(split_docs) >= 2:
            logging.info(f"📝 调试 - Chunk 0 末尾 50 字符: ...{split_docs[0].page_content[-50:]}")
            logging.info(f"📝 调试 - Chunk 1 开头 50 字符: {split_docs[1].page_content[:50]}...")
        
        # 5. 补充元数据（chunk 索引）
        for i, doc in enumerate(split_docs):
            doc.metadata["chunk_index"] = i
        
        # 6. 去重
        file_key = md5(file_path.encode("utf-8")).hexdigest()
        unique_docs = self._remove_duplicate(split_docs, file_key)
        
        logging.info(f"📊 文档处理完成: {len(raw_docs)} 页 -> {len(unique_docs)} 个 chunks")
        
        return unique_docs

    def _clean_docs_light(self, docs: List[Document]) -> List[Document]:
        """轻量清洗，保留 Markdown 结构字符"""
        for doc in docs:
            text = doc.page_content
            # 仅移除不可见字符，保留表格符号 | -
            text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
            # 修复 Markdown 标题格式: "#标题" -> "# 标题" (确保 # 后有空格)
            text = re.sub(r'^(#{1,6})([^\s#])', r'\1 \2', text, flags=re.MULTILINE)
            doc.page_content = text.strip()
        return docs

    def _remove_duplicate(self, docs: List[Document], file_key: str) -> List[Document]:
        """去重逻辑 (保持不变)"""
        if not self.config.enable_duplicate_removal:
            return docs
        
        if file_key not in self.duplicate_cache:
            self.duplicate_cache[file_key] = set()
        
        unique_docs = []
        for doc in docs:
            # 使用内容 hash 去重
            text_hash = md5(doc.page_content.encode("utf-8")).hexdigest()
            if text_hash not in self.duplicate_cache[file_key]:
                self.duplicate_cache[file_key].add(text_hash)
                unique_docs.append(doc)
        
        return unique_docs

    def _parse_pdf_basic(self, file_path: str) -> List[str]:
        """
        基本 PDF 解析（当 pdfplumber 不可用时使用）
        尝试使用 pypdf，如果不可用则返回空列表
        """
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
            return pages
        except ImportError:
            logging.warning("No PDF library installed. Install pdfplumber or pypdf for PDF support: pip install pypdf")
            return []
        except Exception as e:
            logging.error(f"Error loading {file_path}: {e}")
            return []

__all__ = ["IngestionManager"]

