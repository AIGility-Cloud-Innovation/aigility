# [服务层] 对外暴露的统一入口 (RAGManager)
import shutil
# 第一步：添加这几行代码（解决相对导入问题）
import sys
import os
os.environ["TOKENIZERS_PARALLELISM"] = "false"
# 获取当前文件的父目录（aigility/rag）
current_dir = os.path.dirname(os.path.abspath(__file__))
# 获取项目根目录（aigility-master）
root_dir = os.path.dirname(os.path.dirname(current_dir))
# 将根目录添加到Python路径
sys.path.insert(0, root_dir)

# 第二步：修改相对导入为绝对导入（或保留相对导入，此时已识别父包）
from aigility.rag.config import RAGConfig
from aigility.rag.embeddings.factory import EmbeddingFactory
from aigility.rag.vector_stores.factory import VectorStoreFactory
from aigility.rag.ingestion import IngestionManager
import hashlib
from sklearn.feature_extraction.text import TfidfVectorizer
import jieba
import docx
import pdfplumber

class RAGService:
    def __init__(self, config: RAGConfig = None):
        # 如果用户没传配置，使用默认配置
        self.config = config or RAGConfig()
        
        print(f"🔧 Initializing RAG with: Embedding={self.config.embedding.provider}, Store={self.config.vector_store.provider}")

        # 1. 工厂生产 Embedding
        self.embedding_model = EmbeddingFactory.get_embedding_model(self.config.embedding)
        
        # 2. 工厂生产 Vector Store (注入 embedding)
        self.vector_store = VectorStoreFactory.get_vector_store(
            self.config.vector_store, 
            self.embedding_model
        )
        
        # 3. 初始化数据处理模块
        self.ingestion = IngestionManager(self.config.ingestion)
        self.doc_meta_info = {}  # 存储单个文档元信息：{文档名: {"name": "", "keywords": [], "summary": ""}}
        self.global_doc_keywords = []  # 所有文档的核心关键词合集（去重），用于后续主题匹配

    def _read_file(self, file_path: str) -> str:
        """
        统一读取txt/docx/pdf格式文件，提取纯文本内容
        :param file_path: 文件绝对/相对路径
        :return: 清洗后的纯文本字符串（失败返回空字符串）
        """
        # 获取文件后缀（小写）
        file_suffix = os.path.splitext(file_path)[-1].lower()
        pure_text = ""

        try:
            if file_suffix == ".txt":
                # 读取TXT文件（支持utf-8/gbk编码）
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        pure_text = f.read()
                except UnicodeDecodeError:
                    with open(file_path, "r", encoding="gbk") as f:
                        pure_text = f.read()

            elif file_suffix == ".docx":
                # 读取DOCX文件（正文+表格）
                doc = docx.Document(file_path)
                # 提取段落文本
                for para in doc.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        pure_text += para_text + "\n"
                # 提取表格文本（可选，根据你的需求开启）
                for table in doc.tables:
                    for row in table.rows:
                        row_text = "\t".join([cell.text.strip() for cell in row.cells])
                        if row_text:
                            pure_text += row_text + "\n"

            elif file_suffix == ".pdf":
                # 读取PDF文件（仅支持可复制文字的PDF，扫描件无效）
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pure_text += page_text.strip() + "\n"

            else:
                raise ValueError(f"不支持的文件格式：{file_suffix}，仅支持txt/docx/pdf")

            # 文本清洗：去除多余换行、空格、制表符
            pure_text = pure_text.replace("\r", "").replace("\t", " ").replace("  ", " ").strip()
            return pure_text

        except Exception as e:
            print(f"❌ 读取文件{file_path}失败：{str(e)}")
            return ""
    def _extract_doc_meta(self, doc_text: str, doc_name: str) -> dict:
            """
            从纯文本中提取文档元信息（关键词+摘要）
            :param doc_text: 文档纯文本内容
            :param doc_name: 文档文件名（如test.pdf）
            :return: 包含关键词和摘要的元信息字典
            """
            # 1. 生成文档摘要（前200字，避免过长）
            doc_summary = doc_text[:200].strip()
            if len(doc_text) <= 200:
                doc_summary = doc_text.strip()

            # 2. 生成核心关键词（jieba分词 + TF-IDF）
            doc_keywords = []
            try:
                # 第一步：中文分词，用空格分隔词语（符合TF-IDF输入要求）
                word_list = jieba.lcut(doc_text.replace("\n", "").strip())  # 分词
                seg_text = " ".join(word_list)  # 拼接为空格分隔的字符串
                corpus = [seg_text]  # TF-IDF要求输入为文本列表

                # 第二步：TF-IDF提取Top10关键词
                tfidf = TfidfVectorizer(
                    max_features=10,  # 最多提取10个关键词
                    stop_words=None  # 可自定义停用词表（如["的", "是", "在"]）
                )
                tfidf.fit_transform(corpus)  # 训练TF-IDF模型
                doc_keywords = tfidf.get_feature_names_out().tolist()  # 获取关键词列表
            except Exception as e:
                # 兜底方案：若TF-IDF失败，用文件名拆分作为关键词
                print(f"⚠️ 提取{doc_name}关键词失败：{str(e)}，使用文件名兜底")
                doc_keywords = [word for word in doc_name.split(".")[0].split("_") if word.strip()]

            # 3. 组装元信息
            doc_meta = {
                "name": doc_name,
                "keywords": doc_keywords,  # 核心关键词列表
                "summary": doc_summary     # 文档摘要
            }

            # 4. 更新全局元信息（去重关键词）
            self.doc_meta_info[doc_name] = doc_meta
            self.global_doc_keywords.extend(doc_keywords)
            self.global_doc_keywords = list(set(self.global_doc_keywords))  # 去重

            return doc_meta

    # 在 service.py 的 add_file 方法中添加


    def add_file(self, file_path: str):
        """加载文件 -> 切分 -> 存入向量库（带去重）"""
        try:
            if not isinstance(file_path, str):
                raise TypeError(f"file_path must be str, got {type(file_path)}")
            if not os.path.isabs(file_path):
                file_path = os.path.join(current_dir, file_path)
            
            # 计算文件 hash，用于去重
            with open(file_path, "rb") as f:
                file_hash = hashlib.md5(f.read()).hexdigest()
            
            doc_name = os.path.basename(file_path)
            
            # 检查是否已添加过该文件（通过 metadata 查询）
            existing = self.vector_store.get(
                where={"file_hash": file_hash},
                limit=1
            )
            if existing and existing.get("ids"):
                print(f"⚠️ 文件已存在，跳过添加: {file_path}")
                # 但仍然需要加载元信息（如果尚未加载）
                if doc_name not in self.doc_meta_info:
                    doc_text = self._read_file(file_path)
                    if doc_text:
                        self._extract_doc_meta(doc_text, doc_name)
                        print(f"✅ 已加载 {doc_name} 的元信息")
                return
            
            # 文件不存在，需要添加
            print(f"📄 Processing file: {file_path}")
            
            # 读取文件内容
            doc_text = self._read_file(file_path)
            if not doc_text:
                print(f"❌ {doc_name} 提取纯文本为空，无法添加到知识库")
                return
            
            # 生成并存储元信息
            doc_meta = self._extract_doc_meta(doc_text, doc_name)
            print(f"✅ 成功生成 {doc_name} 元信息：关键词={doc_meta['keywords']}")
            
            raw_docs = self.ingestion.load_file(file_path)
            chunks = self.ingestion.process_raw_docs(raw_docs, file_path)
            
            # 给每个 chunk 添加 file_hash 元数据
            for chunk in chunks:
                chunk.metadata["file_hash"] = file_hash
            
            if chunks:
                self.vector_store.add_documents(chunks)
                print(f"✅ Successfully added {len(chunks)} chunks to {self.config.vector_store.provider}.")
            else:
                print("⚠️ No content found in file.")
        except Exception as e:
            print(f"❌ Error adding file {file_path}: {str(e)}")
            raise e

    def search(self, query: str) -> str:
        """检索逻辑"""
        try:
            docs = self.vector_store.similarity_search(query, k=self.config.search_top_k)
            
            if not docs:
                return ""

            # 格式化上下文
            results = []
            for doc in docs:
                source = doc.metadata.get("source", "Unknown")
                content = doc.page_content.replace("\n", " ")
                results.append(f"Source: {source}\nContent: {content}")
            
            return "\n\n".join(results)
        except Exception as e:
            print(f"❌ Search failed: {str(e)}")
            return ""
    def clear_knowledge_base(self):
        """(危险操作) 清空知识库"""
        try:
            if self.config.vector_store.provider == "chroma":
                # Chroma：删除持久化目录
                if os.path.exists(self.config.vector_store.persist_path):
                    shutil.rmtree(self.config.vector_store.persist_path)
                    os.makedirs(self.config.vector_store.persist_path, exist_ok=True)
                    print(f"✅ Chroma知识库 {self.config.vector_store.persist_path} 已清空")
            elif self.config.vector_store.provider == "milvus":
                # Milvus：删除集合
                from pymilvus import utility
                if utility.has_collection(self.config.vector_store.collection_name):
                    utility.drop_collection(self.config.vector_store.collection_name)
                    print(f"✅ Milvus集合 {self.config.vector_store.collection_name} 已删除")
            elif self.config.vector_store.provider == "faiss":  # 新增FAISS清空逻辑
                # FAISS：删除索引文件
                faiss_index_path = os.path.join(
                    self.config.vector_store.persist_path,
                    f"{self.config.vector_store.collection_name}.index"
                )
                if os.path.exists(faiss_index_path):
                    os.remove(faiss_index_path)
                    # 可选：删除整个FAISS目录
                    if os.path.exists(self.config.vector_store.persist_path):
                        shutil.rmtree(self.config.vector_store.persist_path)
                        os.makedirs(self.config.vector_store.persist_path, exist_ok=True)
                    print(f"✅ FAISS索引 {faiss_index_path} 已清空")
            else:
                print(f"⚠️ 不支持清空 {self.config.vector_store.provider} 知识库")
        except Exception as e:
            print(f"❌ 清空知识库失败: {str(e)}")
            
            
            
        # 方案一新增：获取全局文档关键词（供rag_decision节点使用）
    def get_global_keywords(self) -> list:
        return self.global_doc_keywords

    # 方案一新增：获取单个文档元信息
    def get_doc_meta(self, doc_name: str) -> dict:
        return self.doc_meta_info.get(doc_name, {})

    # 方案一新增：获取所有文档元信息
    def get_all_doc_meta(self) -> dict:
        return self.doc_meta_info

# ====================== 测试代码（直接运行） ======================
if __name__ == "__main__":

    # 示例1：使用默认配置（DashScope + Chroma）
    # ------------------------------
    '''
    service = RAGService()
    service.add_file("./test.pdf")
    print(service.search("“无领导小组讨论”面试考察什么"))
    '''
    '''
    输出：
        🔧 Initializing RAG with: Embedding=huggingface, Store=chroma
        🔍 检测到Embedding维度：512
        📄 Processing file: /Users/qyd/Downloads/Edge/aigility-master/aigility/rag/./test.pdf
        ✅ Successfully added 7 chunks to chroma.
        Source: /Users/qyd/Downloads/Edge/aigility-master/aigility/rag/./test.pdf
        Content: 【TXT-PAGE0-CHUNK7】【标题：4. 休假流程：所有休假均需通过企业微信】 4. 休假流程：所有休假均需通过企业微信提交书面申请，附相关证明材料（病假需医疗证明、婚假需结婚证等），按审批权限报批：试用期员工及普通员工由直属领导、部门负责人审批；主管及以上人员需额外经人力资源部、 总经理审批，审批通过后方可休假，未审批擅自休假视为旷工。 5. 其他说明：休假期间员工需保持通讯畅通，紧急工作需配合处理；休假结束后1个工作日内到人力资源部办理销假手续，未按时销假按旷工处理。                                                                                           

        Source: /Users/qyd/Downloads/Edge/aigility-master/aigility/rag/./test.pdf
        Content: 【TXT-PAGE0-CHUNK6】 - 病假：员工凭县级及以上医院出具的诊断证明、病历等材料申请，病假期间工资按公司薪酬制度执行（连续病假不满3天按正常工资80%发放，3天及以上按当地最低工资标准的80%发放），月累计病假超过10天取消当月全勤奖。 - 事假：员工因个人事务需请假的，事假期间无工资 ，月累计事假不超过3天，年累计事假不超过15天，超期按旷工处理。 - 婚假：员工结婚可享受婚假3天，符合晚婚条件（男满25周岁、女满23周岁）的额外增加婚假7天，婚假期间按正常工资发放，需一次性休完，提前15天提交申请。 - 产假：女职工生育享受产假98天，其中产前可休假15天；难产的增加产假15天 ；生育多胞胎的，每多生育1个婴儿增加产假15天，产假期间按当地生育保险相关规定发放生育津贴。 - 丧假：员工直系亲属（父母、配偶、子女）去世，可 享受丧假3天；祖父母、外祖父母、岳父母、公婆去世，可享受丧假1天，丧假期间按正常工资发放。                                                   

        Source: /Users/qyd/Downloads/Edge/aigility-master/aigility/rag/./test.pdf
        Content: 【TXT-PAGE0-CHUNK5】【标题：5. 使用要求：员工需合理使用办公用品，】 5. 使用要求：员工需合理使用办公用品，杜绝浪费，损坏或丢失办公用品需按成本价赔偿；离职时需交回未使用完的办公用品及借用的办公设备（如电脑、打印机等），经人力资源部验收合格后办理离职手续。 五、员工休假管理制度（最新） 1. 目的：保障员工休息休假权利，规范休假管理，平衡工作与生活，提升员工幸福感。 2. 适用范围：公司全体在职员工（含试用期员工、 劳务派遣员工，试用期员工不享受年假，其他休假按规定执行）。 3. 核心休假类型及标准： - 带薪年休假：员工连续工作满1年不满10年的，年休假5天；满10年不满20年的，年休假10天；满20年的，年休假15天。年休假可分段申请，原则上当年休完，未休完部分不结转至次年，特殊情况经总经理审批可结转不超 过5天至次年第一季度。 - 法定节假日：按《中华人民共和国劳动法》规定执行，包括元旦1天、春节3天、劳动节1天、国庆节3天等，具体放假安排以公司年 度通知为准，法定节假日加班按3倍日工资支付报酬。                                                                                            
    

    # ------------------------------
    # 示例2：切换为 DashScope + FAISS
    # ------------------------------
    from aigility.rag.config import VectorStoreConfig
    from aigility.rag.config import EmbeddingConfig
    embedding_config = EmbeddingConfig(
        provider="dashscope",
        model_name="text-embedding-v4",
        api_key=os.getenv("DASHSCOPE_API_KEY")
    )
    vector_store_config = VectorStoreConfig(
        provider="faiss",
        path="./faiss_db"
    )
    config = RAGConfig(embedding=embedding_config, vector_store=vector_store_config)
    service = RAGService(config=config)
    service.add_file("./test.pdf")
    print(service.search("面试考察的能力是什么"))

    # ------------------------------
    # 示例3：切换为 DashScope + Milvus
    # ------------------------------
    
    from aigility.rag.config import VectorStoreConfig
    from aigility.rag.config import EmbeddingConfig
    embedding_config = EmbeddingConfig(
        provider="dashscope",
        model_name="text-embedding-v4",
        api_key=os.getenv("DASHSCOPE_API_KEY")
    )
    vector_store_config = VectorStoreConfig(
        provider="milvus",  # 切换向量库
        url="http://localhost:19530"
    )
    config = RAGConfig(embedding=embedding_config, vector_store=vector_store_config)
    service = RAGService(config=config)
    service.add_file("./test.pdf")
    print(service.search("面试考察的能力是什么"))
    '''
    rag_service = RAGService()

    # 2. 添加文件（自动生成元信息）
    test_file_path = "test.pdf"  # 你的毕业生就业问题大全文件
    rag_service.add_file(test_file_path)

    # 3. 验证元信息
    print("\n" + "="*50)
    print("生成的文档元信息：")
    print("="*50)
    # 3.1 获取所有文档元信息
    all_meta = rag_service.get_all_doc_meta()
    for doc_name, meta in all_meta.items():
        print(f"文档名：{doc_name}")
        print(f"核心关键词：{meta['keywords']}")
        print(f"文档摘要：{meta['summary']}")
        print("-"*30)

    # 3.2 获取全局关键词（供rag_decision节点匹配使用）
    global_keywords = rag_service.get_global_keywords()
    print(f"\n所有文档全局关键词：{global_keywords}")
    