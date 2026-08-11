"""Regression tests for optional-dependency import boundaries."""

from __future__ import annotations

import json
import os
from pathlib import Path
import subprocess
import sys
import textwrap

import pytest


PROJECT_ROOT = Path(__file__).resolve().parents[1]


def _probe_import(statement: str) -> dict:
    probe = textwrap.dedent(
        f"""
        import builtins
        import contextlib
        import io
        import importlib
        import json
        import logging
        import os
        from pathlib import Path
        import sys
        import warnings

        blocked = {{
            "chromadb",
            "dashscope",
            "docx",
            "faiss",
            "jieba",
            "langchain_chroma",
            "langchain_community",
            "langchain_huggingface",
            "langchain_milvus",
            "langchain_text_splitters",
            "lxml",
            "markdown_it",
            "nltk",
            "openpyxl",
            "pandas",
            "pdfplumber",
            "pymilvus",
            "pypdf",
            "qdrant_client",
            "rank_bm25",
            "scipy",
            "sklearn",
            "timem",
            "zai",
        }}
        attempted = []
        real_import = builtins.__import__
        real_import_module = importlib.import_module

        def direct_caller():
            frame = sys._getframe(2)
            return frame.f_globals.get("__name__", "")

        def traced_import(name, globals=None, locals=None, fromlist=(), level=0):
            root = name.split(".", 1)[0]
            importer = (globals or {{}}).get("__name__", "")
            if level == 0 and importer.startswith("aigility") and root in blocked:
                attempted.append({{"importer": importer, "module": name}})
                error = ModuleNotFoundError(f"blocked optional dependency: {{root}}")
                error.name = root
                raise error
            return real_import(name, globals, locals, fromlist, level)

        def traced_import_module(name, package=None):
            root = name.split(".", 1)[0]
            importer = direct_caller()
            if not name.startswith(".") and importer.startswith("aigility") and root in blocked:
                attempted.append({{"importer": importer, "module": name}})
                error = ModuleNotFoundError(f"blocked optional dependency: {{root}}")
                error.name = root
                raise error
            return real_import_module(name, package)

        builtins.__import__ = traced_import
        importlib.import_module = traced_import_module
        before_env = dict(os.environ)
        before_path = list(sys.path)
        captured_stdout = io.StringIO()
        captured_stderr = io.StringIO()
        caught_warnings = []
        warning_events = []
        caught_logs = []
        error = None

        real_warn = warnings.warn

        def traced_warn(
            message,
            category=None,
            stacklevel=1,
            source=None,
            **kwargs,
        ):
            frame = sys._getframe(1)
            warning_category = (
                type(message)
                if isinstance(message, Warning)
                else (category or UserWarning)
            )
            warning_events.append({{
                "category": warning_category.__name__,
                "emitter": frame.f_globals.get("__name__", ""),
                "message": str(message),
            }})
            return real_warn(
                message,
                category=category,
                stacklevel=stacklevel + 1,
                source=source,
                **kwargs,
            )

        warnings.warn = traced_warn

        class CaptureHandler(logging.Handler):
            def emit(self, record):
                caught_logs.append({{
                    "level": record.levelname,
                    "lineno": record.lineno,
                    "message": record.getMessage(),
                    "name": record.name,
                    "pathname": record.pathname,
                }})

        root_logger = logging.getLogger()
        original_level = root_logger.level
        capture_handler = CaptureHandler()
        root_logger.addHandler(capture_handler)
        root_logger.setLevel(logging.DEBUG)

        with contextlib.redirect_stdout(captured_stdout), contextlib.redirect_stderr(captured_stderr):
            with warnings.catch_warnings(record=True) as records:
                warnings.simplefilter("always")
                try:
                    exec({statement!r}, {{}})
                except BaseException as exc:
                    error = f"{{type(exc).__name__}}: {{exc}}"
                unmatched_events = list(warning_events)
                for record in records:
                    category_name = type(record.message).__name__
                    message = str(record.message)
                    emitter = ""
                    for index, event in enumerate(unmatched_events):
                        if (
                            event["category"] == category_name
                            and event["message"] == message
                        ):
                            emitter = event["emitter"]
                            unmatched_events.pop(index)
                            break
                    caught_warnings.append({{
                        "category": category_name,
                        "emitter": emitter,
                        "filename": record.filename,
                        "lineno": record.lineno,
                        "message": message,
                    }})

        warnings.warn = real_warn

        root_logger.removeHandler(capture_handler)
        root_logger.setLevel(original_level)

        package_root = Path(os.environ["AIGILITY_PACKAGE_ROOT"]).resolve()
        aigility_logs = []
        third_party_logs = []
        for record in caught_logs:
            try:
                Path(record["pathname"]).resolve().relative_to(package_root)
            except ValueError:
                emitted_by_aigility = record["name"].startswith("aigility")
            else:
                emitted_by_aigility = True
            if emitted_by_aigility:
                aigility_logs.append(record)
            else:
                third_party_logs.append(record)

        env_changes = {{
            key: [before_env.get(key), os.environ.get(key)]
            for key in set(before_env) | set(os.environ)
            if before_env.get(key) != os.environ.get(key)
        }}
        print(json.dumps({{
            "attempted": attempted,
            "env_changes": env_changes,
            "error": error,
            "logs": aigility_logs,
            "path_changed": before_path != sys.path,
            "stderr": captured_stderr.getvalue(),
            "stdout": captured_stdout.getvalue(),
            "aigility_warnings": [
                warning
                for warning in caught_warnings
                if warning["emitter"].startswith("aigility")
            ],
            "third_party_warnings": [
                warning
                for warning in caught_warnings
                if not warning["emitter"].startswith("aigility")
            ],
            "third_party_logs": third_party_logs,
        }}, ensure_ascii=False))
        """
    )
    env = os.environ.copy()
    env["AIGILITY_PACKAGE_ROOT"] = str(PROJECT_ROOT / "aigility")
    env["PYTHONPATH"] = str(PROJECT_ROOT)
    completed = subprocess.run(
        [sys.executable, "-c", probe],
        cwd=PROJECT_ROOT,
        env=env,
        capture_output=True,
        text=True,
        check=True,
    )
    assert completed.stderr == ""
    return json.loads(completed.stdout)


@pytest.mark.parametrize(
    "statement",
    [
        "import aigility",
        (
            "import sys; import aigility; "
            "assert 'aigility.chat' not in sys.modules; "
            "assert 'aigility.memory' not in sys.modules; "
            "assert 'aigility.rag' not in sys.modules"
        ),
        "from aigility import ADKClient",
        "from aigility.chat import ChatAgent",
        (
            "import sys; from aigility.chat import ChatAgent; "
            "assert 'aigility.chat.service' not in sys.modules; "
            "assert 'aigility.chatflow' not in sys.modules"
        ),
        "from aigility.memory import MemoryConfig",
        (
            "import sys; from aigility.memory import MemoryConfig; "
            "assert 'aigility.memory.providers.timem' not in sys.modules"
        ),
        "from aigility.memory import TimemMemoryProvider",
        "from aigility.memory.providers import TimemMemoryProvider",
        "from aigility.rag import RAGConfig, TimeMRAGClient",
        (
            "import sys; from aigility.rag import RAGConfig, TimeMRAGClient; "
            "assert 'aigility.rag.service' not in sys.modules; "
            "assert 'aigility.rag.ingestion' not in sys.modules"
        ),
        "from aigility.rag import IngestionManager, RAGService",
        "from aigility import *",
        "from aigility.chat import *",
        "from aigility.memory import *",
        "from aigility.rag import *",
        "import aigility.rag.markdown_splitter",
        "import aigility.rag.hybrid_search",
    ],
)
def test_core_public_imports_are_optional_dependency_free_and_side_effect_free(
    statement: str,
):
    result = _probe_import(statement)

    third_party_warnings = result.pop("third_party_warnings")
    third_party_logs = result.pop("third_party_logs")
    for warning in third_party_warnings:
        assert set(warning) == {
            "category",
            "emitter",
            "filename",
            "lineno",
            "message",
        }
        assert warning["emitter"]
        assert not warning["emitter"].startswith("aigility")
    for record in third_party_logs:
        assert not record["name"].startswith("aigility")

    assert result == {
        "aigility_warnings": [],
        "attempted": [],
        "env_changes": {},
        "error": None,
        "logs": [],
        "path_changed": False,
        "stderr": "",
        "stdout": "",
    }


@pytest.mark.parametrize(
    "statement",
    [
        (
            "from aigility.memory import MemoryProviderConfig, TimemMemoryProvider; "
            "TimemMemoryProvider(MemoryProviderConfig(api_key='test-key'))"
        ),
        (
            "from aigility.memory import Memory, MemoryConfig, MemoryProviderConfig; "
            "Memory(config=MemoryConfig(provider=MemoryProviderConfig(api_key='test-key')))"
        ),
    ],
)
def test_enabled_timem_provider_reports_missing_extra_at_first_use(statement: str):
    result = _probe_import(statement)

    assert result["attempted"] == [
        {"importer": "aigility._optional", "module": "timem"}
    ]
    assert "TiMEM memory requires optional dependency 'timem-ai'" in result["error"]
    assert 'pip install "aigility[timem]"' in result["error"]
    assert result["stdout"] == ""
    assert result["stderr"] == ""
    assert result["aigility_warnings"] == []


def _block_optional_import(monkeypatch, module_root: str):
    from aigility import _optional

    real_import_module = _optional.import_module

    def blocked(name: str, package=None):
        if name.split(".", 1)[0] == module_root:
            error = ModuleNotFoundError(f"blocked optional dependency: {module_root}")
            error.name = module_root
            raise error
        return real_import_module(name, package)

    monkeypatch.setattr(_optional, "import_module", blocked)


@pytest.mark.parametrize(
    ("extension", "module_root", "extra"),
    [
        (".csv", "pandas", "doc-excel"),
        (".pdf", "pdfplumber", "doc-pdf"),
        (".docx", "docx", "doc-word"),
    ],
)
def test_document_parser_reports_missing_extra_at_first_use(
    tmp_path: Path,
    monkeypatch,
    extension: str,
    module_root: str,
    extra: str,
):
    from aigility.rag import IngestionConfig, IngestionManager

    _block_optional_import(monkeypatch, module_root)
    path = tmp_path / f"document{extension}"
    path.write_bytes(b"placeholder")
    manager = IngestionManager(IngestionConfig())

    with pytest.raises(ImportError, match=rf"aigility\[{extra}\]"):
        manager.load_file(str(path))


def test_text_splitter_reports_missing_rag_extra_at_first_use(monkeypatch):
    from aigility.rag import IngestionConfig, IngestionManager

    _block_optional_import(monkeypatch, "langchain_text_splitters")
    manager = IngestionManager(IngestionConfig())

    with pytest.raises(ImportError, match=r"aigility\[rag\]"):
        _ = manager.splitter


@pytest.mark.parametrize(
    ("extension", "replacement"), [(".xls", ".xlsx"), (".doc", ".docx")]
)
def test_legacy_document_formats_are_explicitly_unsupported(
    tmp_path: Path,
    extension: str,
    replacement: str,
):
    from aigility.rag import IngestionConfig, IngestionManager

    path = tmp_path / f"legacy{extension}"
    path.write_bytes(b"placeholder")

    with pytest.raises(ValueError, match=rf"convert it to \{replacement}"):
        IngestionManager(IngestionConfig()).load_file(str(path))


@pytest.mark.parametrize(
    "missing_name", ["internal_dependency", "installed_optional.internal"]
)
def test_optional_import_preserves_dependency_internal_import_errors(
    monkeypatch,
    missing_name: str,
):
    from aigility import _optional

    internal_error = ModuleNotFoundError("missing dependency inside installed module")
    internal_error.name = missing_name

    def fail_inside_module(name: str, package=None):
        raise internal_error

    monkeypatch.setattr(_optional, "import_module", fail_inside_module)

    with pytest.raises(ModuleNotFoundError) as caught:
        _optional.import_optional(
            "installed_optional",
            feature="Test feature",
            extra="test-extra",
        )

    assert caught.value is internal_error


def test_explicit_dashscope_embedding_request_reports_its_extra(monkeypatch):
    from aigility.rag import EmbeddingConfig
    from aigility.rag.embeddings import EmbeddingFactory

    monkeypatch.setitem(sys.modules, "dashscope", None)
    monkeypatch.delitem(
        sys.modules,
        "aigility.rag.embeddings.dashscope",
        raising=False,
    )

    with pytest.raises(
        ImportError,
        match=r'pip install "aigility\[embedding-dashscope\]"',
    ):
        EmbeddingFactory.get_embedding_model(
            EmbeddingConfig(
                provider="dashscope",
                model_name="text-embedding-v3",
                api_key="test-key",
            )
        )


def test_explicit_dashscope_rerank_request_reports_its_extra(monkeypatch):
    from aigility.rag import RerankConfig, RerankFactory

    monkeypatch.setitem(sys.modules, "dashscope", None)
    monkeypatch.delitem(sys.modules, "aigility.rag.rerank.dashscope", raising=False)

    with pytest.raises(
        ImportError,
        match=r'pip install "aigility\[rerank-dashscope\]"',
    ):
        RerankFactory.get_reranker(RerankConfig(enabled=True, api_key="test-key"))


def test_dashscope_embedding_preserves_sdk_internal_import_errors(
    tmp_path: Path,
    monkeypatch,
):
    from aigility.rag import EmbeddingConfig
    from aigility.rag.embeddings import EmbeddingFactory

    fake_sdk = tmp_path / "dashscope"
    fake_sdk.mkdir()
    (fake_sdk / "__init__.py").write_text(
        "import dashscope_internal_dependency\n",
        encoding="utf-8",
    )
    monkeypatch.syspath_prepend(str(tmp_path))
    monkeypatch.delitem(sys.modules, "dashscope", raising=False)

    with pytest.raises(ModuleNotFoundError) as caught:
        EmbeddingFactory.get_embedding_model(
            EmbeddingConfig(
                provider="dashscope",
                model_name="text-embedding-v3",
                api_key="test-key",
            )
        )

    assert caught.value.name == "dashscope_internal_dependency"
    assert "aigility[embedding-dashscope]" not in str(caught.value)


def test_huggingface_provider_reports_its_extra_without_mutating_environment(
    monkeypatch,
):
    from aigility.rag import EmbeddingConfig
    from aigility.rag.embeddings.factory import EmbeddingFactory

    _block_optional_import(monkeypatch, "langchain_huggingface")
    before = os.environ.get("TOKENIZERS_PARALLELISM")

    with pytest.raises(ImportError, match=r"aigility\[embedding-huggingface\]"):
        EmbeddingFactory.get_embedding_model(EmbeddingConfig(provider="huggingface"))

    assert os.environ.get("TOKENIZERS_PARALLELISM") == before


def test_enabled_timem_provider_initializes_at_sdk_boundary(monkeypatch):
    from types import SimpleNamespace

    from aigility import _optional
    from aigility.memory import MemoryProviderConfig, TimemMemoryProvider

    created = []

    class FakeAsyncMemory:
        def __init__(self, **kwargs):
            created.append(kwargs)

    real_import_module = _optional.import_module

    def fake_import_module(name: str, package=None):
        if name == "timem":
            return SimpleNamespace(AsyncMemory=FakeAsyncMemory)
        return real_import_module(name, package)

    monkeypatch.setattr(_optional, "import_module", fake_import_module)
    provider = TimemMemoryProvider(
        MemoryProviderConfig(
            api_key="test-key",
            base_url="https://memory.example.test/",
        )
    )

    assert provider.enabled is True
    assert isinstance(provider._client, FakeAsyncMemory)
    assert created == [
        {
            "api_key": "test-key",
            "base_url": "https://memory.example.test",
            "timeout": 90.0,
            "max_retries": 0,
        }
    ]


def test_enabled_memory_propagates_provider_constructor_failure(monkeypatch):
    from types import SimpleNamespace

    from aigility import _optional
    from aigility.memory import Memory, MemoryConfig, MemoryProviderConfig

    provider_error = RuntimeError("provider constructor failed")

    class FailingAsyncMemory:
        def __init__(self, **kwargs):
            raise provider_error

    real_import_module = _optional.import_module

    def fake_import_module(name: str, package=None):
        if name == "timem":
            return SimpleNamespace(AsyncMemory=FailingAsyncMemory)
        return real_import_module(name, package)

    monkeypatch.setattr(_optional, "import_module", fake_import_module)

    with pytest.raises(RuntimeError) as caught:
        Memory(
            config=MemoryConfig(
                provider=MemoryProviderConfig(api_key="test-key")
            )
        )

    assert caught.value is provider_error


@pytest.mark.optional_timem
def test_timem_extra_initializes_without_network_access():
    import timem

    from aigility.memory import MemoryProviderConfig, TimemMemoryProvider

    provider = TimemMemoryProvider(
        MemoryProviderConfig(
            api_key="test-key",
            base_url="https://memory.example.test/",
        )
    )

    assert provider.enabled is True
    assert isinstance(provider._client, timem.AsyncMemory)


@pytest.mark.optional_doc_excel
@pytest.mark.parametrize("extension", [".csv", ".xlsx"])
def test_doc_excel_extra_parses_real_files(tmp_path: Path, extension: str):
    import pandas  # noqa: F401 - a hard assertion that the selected extra is present
    import openpyxl

    from aigility.rag import IngestionConfig, IngestionManager

    path = tmp_path / f"people{extension}"
    if extension == ".csv":
        path.write_text("name,role\nLuna,worker\n", encoding="utf-8")
    else:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.append(["name", "role"])
        sheet.append(["Luna", "worker"])
        workbook.save(path)

    documents = IngestionManager(IngestionConfig()).load_file(str(path))

    assert documents
    assert "name: Luna" in documents[0].page_content
    assert "role: worker" in documents[0].page_content


@pytest.mark.optional_doc_pdf
def test_doc_pdf_extra_parses_a_real_pdf(tmp_path: Path):
    import pdfplumber  # noqa: F401 - the selected extra must be installed
    from reportlab.pdfgen import canvas

    from aigility.rag import IngestionConfig, IngestionManager

    path = tmp_path / "luna.pdf"
    document = canvas.Canvas(str(path))
    document.drawString(72, 720, "Luna PDF boundary")
    document.save()

    parsed = IngestionManager(IngestionConfig()).load_file(str(path))

    assert parsed
    assert "Luna PDF boundary" in parsed[0].page_content


@pytest.mark.optional_doc_pdf
def test_corrupt_pdf_fails_instead_of_returning_an_empty_success(tmp_path: Path):
    import pdfplumber  # noqa: F401 - rule out a missing-extra failure

    from aigility.rag import IngestionConfig, IngestionManager

    path = tmp_path / "corrupt.pdf"
    path.write_bytes(b"%PDF-1.7\nnot-a-valid-document")

    with pytest.raises(Exception) as caught:
        IngestionManager(IngestionConfig()).load_file(str(path))

    assert "aigility[doc-pdf]" not in str(caught.value)


@pytest.mark.optional_doc_word
def test_doc_word_extra_parses_a_real_docx(tmp_path: Path):
    import docx

    from aigility.rag import IngestionConfig, IngestionManager

    path = tmp_path / "luna.docx"
    document = docx.Document()
    document.add_heading("Luna Worker", level=1)
    document.add_paragraph("DOCX boundary content")
    document.save(path)

    parsed = IngestionManager(IngestionConfig()).load_file(str(path))

    assert parsed
    assert "Luna Worker" in parsed[0].page_content
    assert "DOCX boundary content" in parsed[0].page_content


@pytest.mark.optional_doc_markdown
def test_doc_markdown_extra_splits_structured_markdown():
    import markdown_it  # noqa: F401 - the selected extra must be installed

    from aigility.rag.markdown_splitter import MarkdownASTSplitter

    chunks = MarkdownASTSplitter(chunk_size=80, min_chunk_size=1).split_text(
        "# Luna\n\nWorker boundary content."
    )

    assert chunks
    assert "Luna" in chunks[0]
    assert "Worker boundary content" in chunks[0]


@pytest.mark.optional_rag
def test_rag_extra_provides_the_fallback_text_splitter():
    import langchain_text_splitters  # noqa: F401 - hard extra assertion
    import rank_bm25  # noqa: F401 - the accelerator belongs to the RAG extra

    from aigility.rag import IngestionConfig, IngestionManager

    chunks = IngestionManager(
        IngestionConfig(chunk_size=40, chunk_overlap=5)
    ).splitter.split_text("Luna worker boundary. " * 8)

    assert chunks
    assert any("Luna worker boundary" in chunk for chunk in chunks)


def test_enabled_memory_without_credentials_fails_fast(monkeypatch):
    from aigility.memory import Memory, MemoryConfig

    monkeypatch.delenv("TIMEM_API_KEY", raising=False)

    with pytest.raises(ValueError, match="TiMEM memory is enabled"):
        Memory(config=MemoryConfig())


def test_lazy_facades_preserve_all_dir_and_repeated_import_identity():
    import typing

    import aigility
    import aigility.chat as chat
    import aigility.memory as memory
    import aigility.rag as rag

    for module in (aigility, chat, memory, rag):
        assert set(module.__all__).issubset(dir(module))

    first = chat.ChatAgent
    second = chat.ChatAgent
    assert first is second
    assert typing.get_type_hints(first.invoke)["return"].__name__ == "AgentResponse"
